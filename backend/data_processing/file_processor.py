"""
Local file processing module for various document formats.
Supporting PDF, DOCX, TXT, Markdown with encoding detection and error handling.
"""

import logging
import mimetypes
import os
from collections.abc import Generator
from pathlib import Path
from typing import Any, Optional

try:
    import pypdf
except ImportError:
    logging.warning("pypdf not available. PDF processing will be disabled.")
    pypdf = None

try:
    from docx import Document as DocxDocument
except ImportError:
    logging.warning("python-docx not available. DOCX processing will be disabled.")
    DocxDocument = None

try:
    import chardet
except ImportError:
    logging.warning("chardet not available. Using utf-8 encoding by default.")
    chardet = None

from pydantic import BaseModel

logger = logging.getLogger(__name__)


class FileProcessingResult(BaseModel):
    """Result of file processing operation"""
    file_path: str
    content: str
    file_type: str
    success: bool
    error: Optional[str] = None
    metadata: dict[str, Any] = {}


class FileProcessor:
    """
    Local file processor supporting multiple document formats.
    Handles encoding detection, error recovery, and folder traversal.
    """

    # Supported file extensions
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.md': 'markdown',
        '.markdown': 'markdown',
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',  # Limited support
        '.rtf': 'rtf',
        '.html': 'html',
        '.htm': 'html',
        '.py': 'code',
        '.js': 'code',
        '.ts': 'code',
        '.cpp': 'code',
        '.java': 'code',
        '.json': 'json',
        '.xml': 'xml',
        '.csv': 'csv'
    }

    def __init__(self, max_file_size: int = 50 * 1024 * 1024):  # 50MB default
        """
        Initialize file processor.

        Args:
            max_file_size: Maximum file size in bytes (default 50MB)
        """
        self.max_file_size = max_file_size
        logger.info(f"Initialized FileProcessor with max_file_size={max_file_size/1024/1024:.1f}MB")

    def _detect_encoding(self, file_path: str) -> str:
        """Detect file encoding using chardet if available"""
        if not chardet:
            return 'utf-8'

        try:
            with open(file_path, 'rb') as file:
                raw_data = file.read(10000)  # Read first 10KB for detection
                result = chardet.detect(raw_data)
                encoding = result.get('encoding', 'utf-8')
                confidence = result.get('confidence', 0)

                if confidence > 0.7 and encoding:
                    logger.debug(f"Detected encoding {encoding} with confidence {confidence:.2f} for {file_path}")
                    return encoding
                else:
                    logger.debug(f"Low confidence encoding detection for {file_path}, using utf-8")
                    return 'utf-8'
        except Exception as e:
            logger.warning(f"Encoding detection failed for {file_path}: {e}")
            return 'utf-8'

    def _read_text_file(self, file_path: str) -> str:
        """Read plain text file with encoding detection"""
        encoding = self._detect_encoding(file_path)

        # Try detected encoding first, then fallbacks
        encodings_to_try = [encoding, 'utf-8', 'latin-1', 'cp1252']

        for enc in encodings_to_try:
            try:
                with open(file_path, encoding=enc, errors='replace') as file:
                    content = file.read()
                    logger.debug(f"Successfully read {file_path} with encoding {enc}")
                    return content
            except (UnicodeDecodeError, LookupError) as e:
                logger.debug(f"Failed to read {file_path} with encoding {enc}: {e}")
                continue

        raise ValueError(f"Could not read file {file_path} with any supported encoding")

    def _read_pdf_file(self, file_path: str) -> str:
        """Extract text from PDF file"""
        if not pypdf:
            raise ImportError("pypdf library not available for PDF processing")

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)

                if pdf_reader.is_encrypted:
                    logger.warning(f"PDF file {file_path} is encrypted, attempting to decrypt")
                    try:
                        pdf_reader.decrypt("")  # Try empty password
                    except Exception:
                        raise ValueError("Cannot decrypt password-protected PDF") from None

                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1} in {file_path}: {e}")
                        continue

                content = "\n\n".join(text_content)
                logger.info(f"Extracted text from {len(pdf_reader.pages)} pages in {file_path}")
                return content

        except Exception as e:
            logger.error(f"PDF processing failed for {file_path}: {e}")
            raise

    def _read_docx_file(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not DocxDocument:
            raise ImportError("python-docx library not available for DOCX processing")

        try:
            doc = DocxDocument(file_path)

            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            content = "\n\n".join(paragraphs)
            logger.info(f"Extracted text from DOCX file {file_path}")
            return content

        except Exception as e:
            logger.error(f"DOCX processing failed for {file_path}: {e}")
            raise

    def process_file(self, file_path: str) -> FileProcessingResult:
        """
        Process a single file and extract its text content.

        Args:
            file_path: Path to the file to process

        Returns:
            FileProcessingResult with extracted content
        """
        file_path = str(Path(file_path).resolve())

        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return FileProcessingResult(
                    file_path=file_path,
                    content="",
                    file_type="unknown",
                    success=False,
                    error="File not found"
                )

            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                return FileProcessingResult(
                    file_path=file_path,
                    content="",
                    file_type="unknown",
                    success=False,
                    error=f"File too large: {file_size/1024/1024:.1f}MB > {self.max_file_size/1024/1024:.1f}MB"
                )

            # Determine file type
            file_extension = Path(file_path).suffix.lower()
            file_type = self.SUPPORTED_EXTENSIONS.get(file_extension, 'unknown')

            if file_type == 'unknown':
                # Try to guess mime type
                mime_type, _ = mimetypes.guess_type(file_path)
                if mime_type and mime_type.startswith('text/'):
                    file_type = 'text'
                else:
                    return FileProcessingResult(
                        file_path=file_path,
                        content="",
                        file_type=file_type,
                        success=False,
                        error=f"Unsupported file type: {file_extension}"
                    )

            # Extract content based on file type
            content = ""
            if file_type == 'pdf':
                content = self._read_pdf_file(file_path)
            elif file_type == 'docx':
                content = self._read_docx_file(file_path)
            else:
                # Treat as text file (includes markdown, code, etc.)
                content = self._read_text_file(file_path)

            # Basic content validation
            if not content.strip():
                return FileProcessingResult(
                    file_path=file_path,
                    content="",
                    file_type=file_type,
                    success=False,
                    error="No content extracted from file"
                )

            # Metadata
            metadata = {
                "file_size": file_size,
                "file_extension": file_extension,
                "content_length": len(content),
                "mime_type": mimetypes.guess_type(file_path)[0]
            }

            logger.info(f"Successfully processed {file_path} ({file_type}, {len(content)} chars)")

            return FileProcessingResult(
                file_path=file_path,
                content=content,
                file_type=file_type,
                success=True,
                metadata=metadata
            )

        except Exception as e:
            logger.error(f"Failed to process file {file_path}: {e}")
            return FileProcessingResult(
                file_path=file_path,
                content="",
                file_type=file_extension.lstrip('.') if 'file_extension' in locals() else 'unknown',
                success=False,
                error=str(e)
            )

    def process_folder(self, folder_path_str: str, recursive: bool = True) -> Generator[FileProcessingResult, None, None]:
        """
        Process all supported files in a folder.

        Args:
            folder_path: Path to the folder
            recursive: Whether to process subfolders

        Yields:
            FileProcessingResult for each file processed
        """
        folder_path = Path(folder_path_str).resolve()

        if not folder_path.exists():
            yield FileProcessingResult(
                file_path=str(folder_path),
                content="",
                file_type="unknown",
                success=False,
                error="Folder not found"
            )
            return

        if not folder_path.is_dir():
            yield FileProcessingResult(
                file_path=str(folder_path),
                content="",
                file_type="unknown",
                success=False,
                error="Path is not a directory"
            )
            return

        # Get all files
        if recursive:
            file_pattern = "**/*"
        else:
            file_pattern = "*"

        files_found = 0
        files_processed = 0

        for file_path in folder_path.glob(file_pattern):
            if file_path.is_file():
                files_found += 1

                # Check if supported
                if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                    result = self.process_file(str(file_path))
                    files_processed += 1
                    yield result
                else:
                    logger.debug(f"Skipping unsupported file: {file_path}")

        logger.info(f"Folder processing complete: {files_processed}/{files_found} files processed from {folder_path}")

    def get_supported_extensions(self) -> list[str]:
        """Get list of supported file extensions"""
        return list(self.SUPPORTED_EXTENSIONS.keys())

    def is_supported_file(self, file_path: str) -> bool:
        """Check if a file is supported for processing"""
        extension = Path(file_path).suffix.lower()
        return extension in self.SUPPORTED_EXTENSIONS


# Convenience function for creating processor instance
def create_file_processor(config=None) -> FileProcessor:
    """Create and return a FileProcessor instance"""
    if config:
        max_file_size = int(config.max_file_size_mb * 1024 * 1024)
    else:
        max_file_size = 50 * 1024 * 1024  # Default 50MB
    return FileProcessor(max_file_size=max_file_size)
